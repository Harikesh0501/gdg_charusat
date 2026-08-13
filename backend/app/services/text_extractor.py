import io
import re
import logging
import pypdf
import docx

logger = logging.getLogger(__name__)


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Parses resume binary stream in-memory without saving to disk/cloud storage.
    Supports .pdf (via pypdf) and .docx (via python-docx).
    """
    ext = filename.lower().split(".")[-1]
    extracted_text = ""

    if ext == "pdf":
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
            extracted_text = "\n".join(pages_text)
        except Exception as e:
            logger.error(f"PDF extraction failed for {filename}: {e}")
            raise ValueError(f"Could not parse PDF file: {e}")

    elif ext in ["docx", "doc"]:
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            extracted_text = "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction failed for {filename}: {e}")
            raise ValueError(f"Could not parse DOCX file: {e}")

    else:
        raise ValueError(f"Unsupported file extension: .{ext}. Supported formats are .pdf and .docx.")

    # Normalize whitespace
    cleaned_text = re.sub(r"[ \t]+", " ", extracted_text)
    cleaned_text = re.sub(r"\n\s*\n+", "\n", cleaned_text).strip()

    if len(cleaned_text) < 50:
        raise ValueError(
            "Extracted text is too short (< 50 characters). File may be empty or an unparseable scanned image."
        )

    return cleaned_text
